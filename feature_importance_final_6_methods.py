import pandas as pd
import numpy as np
from sklearn.ensemble import RandomForestClassifier, ExtraTreesClassifier, GradientBoostingClassifier
from sklearn.linear_model import LassoCV
from sklearn.feature_selection import RFECV
from sklearn.inspection import permutation_importance
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import matplotlib.pyplot as plt
from sklearn.preprocessing import LabelEncoder

# Load the dataset
data = pd.read_csv('data.csv')

# Drop the 'ID' column
data = data.drop(columns=['ID'])

# Define the target variable
X = data.drop(columns=['class'])
y = data['class']

# Encode the 'class' column into numeric values
label_encoder = LabelEncoder()
y = label_encoder.fit_transform(data['class'])

# Initialize models
random_forest = RandomForestClassifier()
extra_trees = ExtraTreesClassifier()
gradient_boosting = GradientBoostingClassifier()
lasso = LassoCV(cv=5)
# rfecv = RFECV(estimator=RandomForestClassifier(), cv=5)
permutation_importance_model = RandomForestClassifier()

execution_times = {}

# Fit models
start_time = time.time()
random_forest.fit(X, y)
execution_times['Random Forest'] = time.time() - start_time
print("Random Forest Completed")

start_time = time.time()
extra_trees.fit(X, y)
execution_times['Extra Trees'] = time.time() - start_time
print("Extra Trees Feature Importance Completed")

start_time = time.time()
gradient_boosting.fit(X, y)
execution_times['Gradient Boosting'] = time.time() - start_time
print("Gradient Boosting Feature Importance Completed")

start_time = time.time()
lasso.fit(X, y)
execution_times['L1 Regularization (Lasso)'] = time.time() - start_time
print("L1 Regularization (Lasso) Completed")

'''
start_time = time.time()
rfecv.fit(X, y)
# Get the selected features from RFECV
selected_features = X.columns[rfecv.support_]
# Train a RandomForestClassifier on the selected features
rfecv_importance_model = RandomForestClassifier()
rfecv_importance_model.fit(X[selected_features], y)

# Calculate feature importances for the selected features
rfecv_importances = rfecv_importance_model.feature_importances_
execution_times['RFECV'] = time.time() - start_time
print("Recursive Feature Elimination with Cross-Validation (RFECV) Completed")
'''

start_time = time.time()
permutation_importance_model.fit(X, y)
permutation_importance_result = permutation_importance(permutation_importance_model, X, y, n_repeats=30, random_state=42)
execution_times['Permutation Importance'] = time.time() - start_time
print("Permutation Importance Completed")

feature_names = X.columns.tolist()  # Reset to all features initially
feature_importance_df = pd.DataFrame(columns=['Method'] + feature_names)

# Calculate and store feature importance scores for each method
feature_importance_df.loc[0] = ['Random Forest'] + list(random_forest.feature_importances_)
feature_importance_df.loc[1] = ['Extra Trees'] + list(extra_trees.feature_importances_)
feature_importance_df.loc[2] = ['Gradient Boosting'] + list(gradient_boosting.feature_importances_)
feature_importance_df.loc[3] = ['L1 Regularization (Lasso)'] + list(np.abs(lasso.coef_))
# feature_importance_df.loc[4] = ['RFECV'] + list(rfecv_importances)  # Use rfecv_importances obtained earlier
feature_importance_df.loc[5] = ['Permutation Importance'] + list(permutation_importance_result.importances_mean)

# # Ensure 'Method' column is added
# feature_importance_df.insert(0, 'Method', ['Random Forest', 'Extra Trees', 'Gradient Boosting', 'L1 Regularization (Lasso)', 'Permutation Importance'])

# # Rank the features based on importance scores (descending order)
# feature_importance_df = feature_importance_df.sort_values(by='Method', ascending=False, axis=1)

# Save the ranked feature importance DataFrame to a CSV file
feature_importance_df.to_csv('feature_importance.csv', index=False)

# Find the common features which obtained zero importance across all the methods
# zero_importance_features = feature_importance_df.columns[feature_importance_df.iloc[:, 1:].apply(lambda x: all(x == 0), axis=0)].tolist()
zero_importance_features = feature_importance_df.columns[1:][feature_importance_df.iloc[:, 1:].apply(lambda x: all(x == 0), axis=0)].tolist()
# Print the common features with zero importance
print("Features with zero importance across all methods:", zero_importance_features)


# # Rank the features based on importance scores (descending order)
# feature_importance_df = feature_importance_df.sort_values(by='Method', ascending=False, axis=1)

""" # Create a Word document
doc = Document()

# Add a title to the document
doc.add_heading('Feature Importance Results', level=1)

# Loop through each method and add the results to the document
for _, row in feature_importance_df.iterrows():
    method_name = row['Method']
    doc.add_heading(method_name, level=2)
    
    # Create a table to display the results
    table = doc.add_table(rows=1, cols=len(feature_names) + 1)
    table.autofit = True
    
    # Add headers to the table
    headers = table.rows[0].cells
    headers[0].text = 'Feature'
    for i, feature in enumerate(feature_names):
        headers[i + 1].text = feature
    
    # Add data rows to the table
    for i, feature in enumerate(row[1:]):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = str(feature)
    
    # Add space between methods
    doc.add_paragraph()
    
# Save the Word document
doc.save('feature_importance_results.docx')

print("Feature importance results have been saved to 'feature_importance_results.docx'")
 """
# Plot execution times
methods = ['Random Forest', 'Extra Trees', 'Gradient Boosting',  'L1 Regularization (Lasso)', 'Permutation Importance']#'RFECV', 'Permutation Importance']
times = [execution_times[method] for method in methods]

plt.figure(figsize=(10, 6))
plt.barh(methods, times)
plt.xlabel("Execution Time (seconds)")
plt.ylabel("Feature Importance Method")
plt.title("Execution Time of Feature Importance Methods")
plt.tight_layout()
plt.show()

